# views.py
from django.views import View
from django.views.generic import ListView
from django.db import connection
from django.http import HttpResponse
import docx
from io import BytesIO
from .forms import SubjectDetailsForm


class SubjectDetailsListView(ListView):
    template_name = 'subjectdetails/subject_details_list.html'
    context_object_name = 'subject_details'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        form = SubjectDetailsForm()  # Assuming you have a form for something else
        context['form'] = form
        # Fetch data from the SubjectDetails view using raw SQL query
        with connection.cursor() as cursor:
            cursor.execute("SELECT * FROM SubjectDetails")
            rows = cursor.fetchall()
            subject_details = []
            for row in rows:
                subject_detail = {
                    'study_cycle': row[0],
                    'subject': row[1],
                    'education_form': row[2],
                    'number_of_students': row[3],
                    'textbooks': row[4].split(', ') if row[4] else [],  # Split textbooks string into a list
                    'web_materials': row[5].split(', ') if row[5] else [],  # Split web_materials string into a list
                }
                subject_details.append(subject_detail)
            context['subject_details'] = subject_details
        return context

    def get_queryset(self):
        # Return an empty queryset
        return []

class DownloadDocxView(View):
    def get(self, request, *args, **kwargs):
        # Fetch data from the SubjectDetails view using raw SQL query
        with connection.cursor() as cursor:
            cursor.execute("SELECT * FROM SubjectDetails")
            rows = cursor.fetchall()
            subject_details = []
            for row in rows:
                subject_detail = {
                    'study_cycle': row[0],
                    'subject': row[1],
                    'education_form': row[2],
                    'number_of_students': row[3],
                    'textbooks': row[4].split(', ') if row[4] else [],
                    'web_materials': row[5].split(', ') if row[5] else [],
                }
                subject_details.append(subject_detail)

            # Create a new Word document
            doc = docx.Document()

            # Set the page orientation to landscape
            section = doc.sections[-1]
            section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height

            # Add the header content
            header = doc.add_paragraph("Форма 5\nСВЕДЕНИЯ\nоб учебно-методическом обеспечении образовательной деятельности\n"
                                       "Кыргызский государственный технический университет им.И.Раззакова\n"
                                       "(название юридического лица)\n"
                                       "710200 – Информационные системы и технологии\n"
                                       "(название образовательной программы)", style='Title')
            header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add the table to the document
            table = doc.add_table(rows=1, cols=6)
            hdr_cells = table.rows[0].cells
            headers = ['Study Cycle', 'Subject', 'Education Form', 'Number of Students', 'Textbooks', 'Web Materials']
            for i, header in enumerate(headers):
                cell = hdr_cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.size = Pt(12)
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            for detail in subject_details:
                row_cells = table.add_row().cells
                row_cells[0].text = detail['study_cycle']
                row_cells[1].text = detail['subject']
                row_cells[2].text = detail['education_form']
                row_cells[3].text = str(detail['number_of_students'])
                row_cells[4].text = ', '.join(detail['textbooks']) if detail['textbooks'] else ''
                row_cells[5].text = ', '.join(detail['web_materials']) if detail['web_materials'] else ''
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(12)

            # Save the document to a BytesIO object
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Create the HttpResponse object with the appropriate MIME type
            response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=subject_details.docx'
            return response


